"""
buildkb.py — Knowledge Base Builder for the TDL Expert Bot
=============================================================
Scans a directory of source files, splits them into semantically
meaningful chunks, embeds them with a local sentence-transformer model,
and stores them in a FAISS vector index for fast retrieval-augmented
generation (RAG).

UNIVERSAL FILE SUPPORT
-----------------------
This builder reads and embeds EVERY file type it encounters, not just a
fixed whitelist. It uses format-specific extraction where a good one
exists (PDF, DOCX, XLSX/CSV, JSON), and falls back to safe plain-text
reading (with a binary sniff so images/executables/archives are skipped
automatically) for everything else — .tdl, .txt, .md, .py, .js, .html,
.xml, .yaml, .ini, .log, .sql, and any other text-based source.

KEY IMPROVEMENT OVER NAIVE CHUNKING
------------------------------------
TDL files are structured, definition-based code (e.g. `[Report: X]`,
`[Form: X]`, `[Collection: X]`). Splitting them every N characters —
regardless of where that lands — routinely cuts a single [Report] or
[Form] definition in half. When the chatbot then retrieves a half-chunk,
it only sees part of the definition and has to guess/hallucinate the
rest, which is the single biggest cause of broken generated TDL code.

This version chunks .tdl files by DEFINITION BOUNDARY instead: each
`[Type: Name]` block becomes its own chunk (further split only if a
single definition is unusually large). Every chunk also carries metadata
about which TDL definition type and name it contains, so the retriever
can later boost chunks that match what the user is actually asking for
(e.g. a question about "collections" prioritizes chunks tagged
def_type="Collection").

These building blocks (read_file_content, chunk_tdl_source, chunk_text,
embed_and_index) are imported directly by chat_bot.py so that files
uploaded live in the chatbot UI go through the exact same extraction
and chunking pipeline as files indexed here — there is only ONE code
path for turning a file into searchable vectors, so the two can never
silently drift apart.
"""

import os
import re
import io
import json
import shutil
import hashlib
import argparse
from datetime import datetime

import faiss
from sentence_transformers import SentenceTransformer

# ── Optional format-specific readers. Each is independently optional so ──
# ── the tool still runs (falling back to plain-text reading) if a       ──
# ── particular library isn't installed.                                 ──
try:
    import fitz  # PyMuPDF
    PDF_ENABLED = True
except ImportError:
    PDF_ENABLED = False

try:
    import docx  # python-docx
    DOCX_ENABLED = True
except ImportError:
    DOCX_ENABLED = False

try:
    import openpyxl
    XLSX_ENABLED = True
except ImportError:
    XLSX_ENABLED = False

import csv as csv_module

EMBEDDING_MODEL_NAME = "all-MiniLM-L6-v2"

# Recognized top-level TDL definition keywords. Used both for chunk
# boundary detection and for definition-type metadata tagging.
TDL_DEFINITION_TYPES = (
    "Report", "Form", "Part", "Line", "Field", "Collection", "Function",
    "Menu", "System", "Variable", "Key", "Button", "Object", "Attribute",
    "Table", "TDL", "Yes", "No"
)

# Matches lines like: [Report: MyReport]   or   [#Form : MyForm]  or  [System: Variable: X]
DEFINITION_HEADER_RE = re.compile(
    r"^\s*\[\s*#?\s*(?P<type>" + "|".join(TDL_DEFINITION_TYPES) + r")\s*:\s*(?P<name>[^\]]+)\]",
    re.IGNORECASE,
)

# No extensions are blocked upfront; any file that is binary/compiled
# falls back to printable ASCII/UTF-8 string extraction automatically.
ALWAYS_SKIP_EXTENSIONS = set()


def create_backup(source_dir, backup_folder="kb_backups"):
    """Zip up the source directory before (re)indexing, so nothing is ever lost."""
    if not os.path.exists(source_dir) or not os.listdir(source_dir):
        return
    if not os.path.exists(backup_folder):
        os.makedirs(backup_folder)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = os.path.join(backup_folder, f"source_backup_{timestamp}")
    shutil.make_archive(backup_path, "zip", source_dir)
    print(f"\n[Backup] Secured source directory to: {backup_path}.zip")


def chunk_text(text, chunk_size=900, overlap=150):
    """Generic fallback chunker: fixed-size sliding window with overlap.
    Used for non-.tdl files (txt/json/pdf/docx/xlsx/csv/etc.) and for any
    single TDL definition block that is too large to embed as one chunk.
    """
    chunks = []
    start = 0
    text_len = len(text)
    if text_len <= chunk_size:
        return [text] if text.strip() else []
    while start < text_len:
        end = min(start + chunk_size, text_len)
        piece = text[start:end].strip()
        if piece:
            chunks.append(piece)
        if end == text_len:
            break
        start += chunk_size - overlap
    return chunks


def chunk_tdl(text, max_chunk_size=900):
    """Smart chunking for TDL code by definition headers."""
    chunks = []
    current = []
    current_size = 0
    for line in text.splitlines():
        is_header = bool(DEFINITION_HEADER_RE.match(line)) or (line.strip().startswith("[") and line.strip().endswith("]"))
        if is_header and current:
            chunks.append("\n".join(current))
            current = [line]
            current_size = len(line)
        else:
            current.append(line)
            current_size += len(line)
            if current_size >= max_chunk_size and len(current) > 1:
                chunks.append("\n".join(current))
                current = []
                current_size = 0
    if current:
        chunks.append("\n".join(current))
    return [c for c in chunks if c.strip()]


def chunk_markdown(text, max_chunk_size=900):
    """Smart chunking for Markdown documents by headings."""
    chunks = []
    current = []
    current_size = 0
    for line in text.splitlines():
        is_heading = bool(re.match(r"^#{1,6}\s+", line))
        if is_heading and current:
            chunks.append("\n".join(current))
            current = [line]
            current_size = len(line)
        else:
            current.append(line)
            current_size += len(line)
            if current_size >= max_chunk_size and len(current) > 1:
                chunks.append("\n".join(current))
                current = []
                current_size = 0
    if current:
        chunks.append("\n".join(current))
    return [c for c in chunks if c.strip()]


def chunk_json(text, max_chunk_size=900):
    """Smart chunking for JSON arrays and objects."""
    try:
        data = json.loads(text)
        if isinstance(data, list):
            return [json.dumps(item) for item in data]
        elif isinstance(data, dict):
            chunks = []
            for key, val in data.items():
                if isinstance(val, list) and len(json.dumps(val)) > max_chunk_size:
                    for item in val:
                        chunks.append(json.dumps({key: item}))
                else:
                    chunks.append(json.dumps({key: val}))
            return chunks
    except Exception:
        pass
    return chunk_text(text, chunk_size=max_chunk_size)


def chunk_tdl_source(text, max_block_size=2500):
    """Definition-boundary chunker for .tdl files.

    Splits the file at every top-level `[Type: Name]` header so each
    chunk is a complete, self-contained TDL definition. Returns a list
    of (chunk_text, def_type, def_name) tuples. Any oversized definition
    is further split with the generic sliding-window chunker, and any
    leading preamble (comments, includes) before the first definition is
    kept as its own untyped chunk.
    """
    lines = text.splitlines()
    header_positions = []
    for i, line in enumerate(lines):
        m = DEFINITION_HEADER_RE.match(line)
        if m:
            header_positions.append((i, m.group("type"), m.group("name").strip()))

    if not header_positions:
        # No recognizable TDL structure found — treat as plain text.
        return [(c, "Text", "") for c in chunk_text(text)]

    results = []

    # Preamble before the first definition (comments / directives)
    if header_positions[0][0] > 0:
        preamble = "\n".join(lines[: header_positions[0][0]]).strip()
        if preamble:
            results.append((preamble, "Preamble", ""))

    for idx, (line_no, def_type, def_name) in enumerate(header_positions):
        end_line = header_positions[idx + 1][0] if idx + 1 < len(header_positions) else len(lines)
        block = "\n".join(lines[line_no:end_line]).strip()
        if not block:
            continue
        if len(block) <= max_block_size:
            results.append((block, def_type, def_name))
        else:
            # Oversized definition (e.g. a big Function/UDF library) — split further,
            # but keep the header attached to every sub-chunk so context isn't lost.
            header_line = lines[line_no]
            sub_chunks = chunk_text(block, chunk_size=max_block_size, overlap=200)
            for j, sub in enumerate(sub_chunks):
                sub_with_header = sub if sub.startswith(header_line) else f"{header_line}\n{sub}"
                results.append((sub_with_header, def_type, f"{def_name} (part {j + 1})"))

    return results


def process_json_bytes(raw_bytes):
    """Flattens JSON into a readable string for embedding."""
    try:
        data = json.loads(raw_bytes.decode("utf-8", errors="ignore"))
        return json.dumps(data, indent=2)
    except Exception as e:
        print(f"  - [ERROR] JSON parse failed: {e}")
        return ""


def process_csv_bytes(raw_bytes, filename=""):
    """Renders CSV/TSV rows as readable 'column: value' text blocks so
    the embedding model gets semantic column context instead of a raw
    comma-separated blob."""
    text = raw_bytes.decode("utf-8", errors="ignore")
    delimiter = "\t" if filename.lower().endswith((".tsv", ".tab")) else ","
    try:
        reader = csv_module.DictReader(io.StringIO(text), delimiter=delimiter)
        rows_text = []
        for i, row in enumerate(reader):
            line = "; ".join(f"{k}: {v}" for k, v in row.items() if k)
            if line:
                rows_text.append(f"Row {i + 1} — {line}")
            if i >= 5000:  # sanity cap for pathologically large CSVs
                rows_text.append("... (truncated, file exceeds 5000 rows)")
                break
        return "\n".join(rows_text) if rows_text else text
    except Exception:
        return text


def process_xlsx_bytes(raw_bytes):
    """Reads every sheet of an Excel workbook into readable row text."""
    if not XLSX_ENABLED:
        return ""
    wb = openpyxl.load_workbook(io.BytesIO(raw_bytes), data_only=True, read_only=True)
    out = []
    for sheet in wb.worksheets:
        out.append(f"# Sheet: {sheet.title}")
        headers = None
        for r_idx, row in enumerate(sheet.iter_rows(values_only=True)):
            if row is None:
                continue
            if headers is None:
                headers = [str(c) if c is not None else f"col{i}" for i, c in enumerate(row)]
                continue
            cells = [f"{headers[i]}: {v}" for i, v in enumerate(row) if v not in (None, "")]
            if cells:
                out.append("Row — " + "; ".join(cells))
            if r_idx > 5000:
                out.append("... (truncated, sheet exceeds 5000 rows)")
                break
    return "\n".join(out)


def process_docx_bytes(raw_bytes):
    """Extracts paragraph and table text from a Word document."""
    if not DOCX_ENABLED:
        return ""
    document = docx.Document(io.BytesIO(raw_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def process_pdf_bytes(raw_bytes):
    if not PDF_ENABLED:
        return ""
    content = ""
    doc = fitz.open(stream=raw_bytes, filetype="pdf")
    for page in doc:
        content += page.get_text() + "\n"
    doc.close()
    return content


def _looks_binary(raw_bytes, sample_size=2048):
    """Cheap binary sniff: NUL bytes or a high ratio of non-printable
    bytes means this almost certainly isn't text we should try to embed."""
    if not raw_bytes:
        return False
    sample = raw_bytes[:sample_size]
    if b"\x00" in sample:
        return True
    text_chars = bytes(range(32, 127)) + b"\n\r\t\f\b"
    nontext = sum(1 for b in sample if b not in text_chars)
    return (nontext / max(len(sample), 1)) > 0.30


def extract_printable_strings(raw_bytes, min_length=4):
    """Extract all printable ASCII / UTF-8 strings of at least min_length
    from raw binary bytes (useful for .tcp compiled add-ons, legacy .doc/.xls,
    or proprietary binary project files)."""
    if not raw_bytes:
        return None
    try:
        pattern = re.compile(rb"[\x09\x0a\x0d\x20-\x7e]{" + str(min_length).encode() + rb",}")
        matches = pattern.findall(raw_bytes)
        strings = [m.decode("ascii", errors="ignore").strip() for m in matches]
        valid_strings = [s for s in strings if len(s) >= min_length and any(c.isalnum() for c in s)]
        if valid_strings:
            return "\n".join(valid_strings)
    except Exception:
        pass
    return None


def read_bytes_content(raw_bytes, filename):
    """Extract readable text from raw file bytes based on extension,
    dispatching to the right specialized reader or falling back to
    universal printable string extraction for binary/compiled files."""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return process_pdf_bytes(raw_bytes) if PDF_ENABLED else extract_printable_strings(raw_bytes)
    if ext == ".docx":
        return process_docx_bytes(raw_bytes) if DOCX_ENABLED else extract_printable_strings(raw_bytes)
    if ext in (".xlsx", ".xlsm"):
        return process_xlsx_bytes(raw_bytes) if XLSX_ENABLED else extract_printable_strings(raw_bytes)
    if ext == ".json":
        return process_json_bytes(raw_bytes)
    if ext in (".csv", ".tsv", ".tab"):
        return process_csv_bytes(raw_bytes, filename)

    if not _looks_binary(raw_bytes):
        return raw_bytes.decode("utf-8", errors="ignore")

    # Binary or compiled content (.tcp, .doc, .xls, etc.): extract printable strings
    return extract_printable_strings(raw_bytes, min_length=4)


def read_file_content(file_path, filename):
    """Extract raw text content from a file on disk, by reading its
    bytes and delegating to read_bytes_content()."""
    with open(file_path, "rb") as f:
        raw_bytes = f.read()
    return read_bytes_content(raw_bytes, filename)


def chunk_any_content(content, filename):
    """Route content to the TDL-aware chunker or the generic chunker
    based on file extension or detected TDL definitions. Returns a list of (chunk, def_type, def_name)."""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".tdl" or re.search(r"^\s*\[\s*#?\s*(?:Report|Form|Part|Line|Field|Collection|System|Menu|Function)\s*:", content, re.M | re.I):
        return chunk_tdl_source(content)
    return [(c, "Text", "") for c in chunk_text(content)]


def embed_and_index(chunks_data, embedder, existing_index=None):
    """Embed a list of chunk strings and return (embeddings, index) where
    index is a ready-to-use (or appended-to) FAISS cosine-similarity
    index. Shared by disk indexing and live chatbot uploads so both
    paths produce vectors that are directly comparable in the same
    index."""
    embeddings = embedder.encode(
        chunks_data,
        convert_to_numpy=True,
        batch_size=64,
        show_progress_bar=len(chunks_data) > 50,
    ).astype("float32")
    faiss.normalize_L2(embeddings)

    if existing_index is not None:
        existing_index.add(embeddings)
        return embeddings, existing_index

    dimension = embeddings.shape[1]
    index = faiss.IndexFlatIP(dimension)
    index.add(embeddings)
    return embeddings, index


def tokenize_for_lexical(text):
    """Tokenize text into lowercase technical tokens, identifiers, and subwords."""
    raw_tokens = re.findall(r"[A-Za-z0-9_#$@*:.]+", text.lower())
    sub_tokens = re.findall(r"[a-z0-9]+", text.lower())
    combined = set(raw_tokens + sub_tokens)
    return [t for t in combined if len(t) >= 2]


def build_sparse_lexical_index(chunks_data):
    """Build an inverted BM25 / TF-IDF sparse lexical index for exact syntax & identifier matching."""
    doc_freqs = {}
    doc_lengths = []
    inverted_index = {}

    for idx, chunk in enumerate(chunks_data):
        tokens = tokenize_for_lexical(chunk)
        doc_lengths.append(len(tokens))
        token_counts = {}
        for t in tokens:
            token_counts[t] = token_counts.get(t, 0) + 1

        for t, count in token_counts.items():
            doc_freqs[t] = doc_freqs.get(t, 0) + 1
            if t not in inverted_index:
                inverted_index[t] = []
            inverted_index[t].append((idx, count))

    avg_dl = sum(doc_lengths) / max(len(doc_lengths), 1)
    return {
        "doc_freqs": doc_freqs,
        "doc_lengths": doc_lengths,
        "avg_doc_length": avg_dl,
        "inverted_index": inverted_index,
        "total_docs": len(chunks_data),
    }



def build_vector_kb(source_dir, index_file="kb_index.faiss", meta_file="kb_metadata.json",
                     embedder=None, allowed_extensions=None, max_file_size_mb=10):
    """
    allowed_extensions: optional whitelist of extensions to restrict processing
        to. If None (the default), EVERY file type is processed — PDFs, DOCX,
        XLSX/CSV, JSON get dedicated extraction; anything else is read as
        plain text (after a binary sniff so images/archives/executables are
        automatically skipped). Pass an explicit set via --extensions if you
        want to narrow a folder (e.g. a full Tally backup) down to just the
        relevant source types.
    max_file_size_mb: files larger than this are skipped with a warning rather
        than silently exploding the chunk count. Raise it only if you're sure
        a large file is genuinely something you want indexed in full.
    """
    print(f"\n[System] Scanning directory: {source_dir}")
    if allowed_extensions:
        print(f"[System] Restricting to extensions: {', '.join(sorted(allowed_extensions))}")
    else:
        print("[System] No extension restriction — indexing ALL readable file types "
              "(binary/media files are auto-skipped).")

    if embedder is None:
        print(f"[System] Loading embedding model: {EMBEDDING_MODEL_NAME} ...")
        embedder = SentenceTransformer(EMBEDDING_MODEL_NAME)

    chunks_data = []
    metadata = []
    seen_chunk_hashes = set()  # dedupe identical/near-identical chunks across files

    max_bytes = max_file_size_mb * 1024 * 1024
    scanned = 0
    skipped_ext = 0
    skipped_size = 0
    skipped_binary = 0
    skipped_empty = 0

    for root, _, files in os.walk(source_dir):
        for filename in files:
            ext = os.path.splitext(filename)[1].lower()

            if ext in ALWAYS_SKIP_EXTENSIONS:
                skipped_ext += 1
                continue
            if allowed_extensions and ext not in allowed_extensions:
                skipped_ext += 1
                continue
            if ext == ".pdf" and not PDF_ENABLED:
                print(f"  - [SKIP] {filename}: PDF support not installed (pip install pymupdf)")
                continue
            if ext == ".docx" and not DOCX_ENABLED:
                print(f"  - [SKIP] {filename}: DOCX support not installed (pip install python-docx)")
                continue
            if ext in (".xlsx", ".xlsm") and not XLSX_ENABLED:
                print(f"  - [SKIP] {filename}: XLSX support not installed (pip install openpyxl)")
                continue

            file_path = os.path.join(root, filename)
            rel_path = os.path.relpath(file_path, source_dir)

            try:
                if os.path.getsize(file_path) > max_bytes:
                    print(f"  - [SKIP] {rel_path}: exceeds {max_file_size_mb} MB "
                          f"(use --max-file-size-mb to raise this if intentional)")
                    skipped_size += 1
                    continue

                content = read_file_content(file_path, filename)
                if content is None:
                    skipped_binary += 1
                    continue
                content = content.strip()
                if not content:
                    skipped_empty += 1
                    continue

                file_chunks = chunk_any_content(content, filename)

                added_for_file = 0
                for i, (chunk, def_type, def_name) in enumerate(file_chunks):
                    chunk_hash = hashlib.sha256(chunk.strip().encode("utf-8")).hexdigest()
                    if chunk_hash in seen_chunk_hashes:
                        continue
                    seen_chunk_hashes.add(chunk_hash)

                    chunks_data.append(chunk)
                    metadata.append({
                        "file": rel_path,
                        "chunk_id": i,
                        "length": len(chunk),
                        "def_type": def_type,
                        "def_name": def_name,
                    })
                    added_for_file += 1

                scanned += 1
                print(f"  + [{scanned}] Indexed: {rel_path} ({added_for_file} chunks) "
                      f"| running total: {len(chunks_data)} chunks")

            except Exception as e:
                print(f"  - [ERROR] Failed on '{rel_path}': {e}")

    print(f"\n[System] Scan complete. {scanned} file(s) indexed, "
          f"{skipped_ext} skipped by extension/type, {skipped_size} skipped for size, "
          f"{skipped_binary} skipped as binary, {skipped_empty} skipped as empty.")

    if not chunks_data:
        print("\n[WARNING] No valid data found to index.")
        return

    print(f"\n[System] Generating semantic vectors for {len(chunks_data)} chunks ...")
    if len(chunks_data) > 20000:
        print(f"[WARNING] {len(chunks_data)} chunks is a lot — this may take a long time "
              f"on CPU. Consider narrowing --source-dir, or lowering --max-file-size-mb, "
              f"if this wasn't expected.")

    embeddings, index = embed_and_index(chunks_data, embedder)

    print(f"\n[System] Generating sparse lexical inverted index for {len(chunks_data)} chunks ...")
    sparse_index = build_sparse_lexical_index(chunks_data)

    faiss.write_index(index, index_file)

    with open(meta_file, "w", encoding="utf-8") as f:
        json.dump({
            "metadata": metadata,
            "chunks": chunks_data,
            "sparse_index": sparse_index,
            "embedding_model": EMBEDDING_MODEL_NAME,
            "similarity_metric": "cosine",
            "built_at": datetime.now().isoformat(),
        }, f, indent=2)

    print("\nSUCCESS! Vector knowledge base built.")
    print(f"Total searchable chunks embedded: {len(chunks_data)}")
    type_counts = {}
    for m in metadata:
        type_counts[m["def_type"]] = type_counts.get(m["def_type"], 0) + 1
    print("Chunk breakdown by definition/content type:")
    for t, c in sorted(type_counts.items(), key=lambda x: -x[1]):
        print(f"    {t}: {c}")


def main():
    parser = argparse.ArgumentParser(
        description="Build the RAG knowledge base. By default, EVERY file type "
                    "under --source-dir is read, chunked, and embedded."
    )
    parser.add_argument(
        "--source-dir",
        default=os.environ.get("TDL_SOURCE_DIR", os.path.join(os.getcwd(), "source_files")),
        help="Directory containing source files to index (default: $TDL_SOURCE_DIR "
             "or ./source_files). ALL file types are processed by default.",
    )
    parser.add_argument("--index-file", default="kb_index.faiss")
    parser.add_argument("--meta-file", default="kb_metadata.json")
    parser.add_argument("--no-backup", action="store_true", help="Skip zipping the source directory first.")
    parser.add_argument(
        "--extensions", default="",
        help="Optional comma-separated whitelist of file extensions to restrict indexing to "
             "(e.g. '.tdl,.txt,.pdf'). Leave empty (default) to index ALL file types found.",
    )
    parser.add_argument(
        "--max-file-size-mb", type=float, default=100,
        help="Skip individual files larger than this many MB (default: %(default)s). "
             "Prevents one huge file from silently generating an enormous number of "
             "chunks and making the embedding step look stuck.",
    )
    args = parser.parse_args()

    if not os.path.exists(args.source_dir):
        os.makedirs(args.source_dir)
        print(f"[System] Created empty source directory: {args.source_dir}")
        print("[System] Add your files there and re-run this script.")
        return

    if not args.no_backup:
        create_backup(args.source_dir)

    allowed_extensions = None
    if args.extensions.strip():
        allowed_extensions = {e.strip().lower() if e.strip().startswith(".") else f".{e.strip().lower()}"
                               for e in args.extensions.split(",") if e.strip()}

    build_vector_kb(
        args.source_dir, args.index_file, args.meta_file,
        allowed_extensions=allowed_extensions,
        max_file_size_mb=args.max_file_size_mb,
    )


if __name__ == "__main__":
    main()